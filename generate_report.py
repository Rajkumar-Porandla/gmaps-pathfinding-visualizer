import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(val))

def set_cell_shading(cell, color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def create_report_v2():
    doc = docx.Document()
    
    # Page settings: A4 margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Text styles
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    
    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    
    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Caption style
    caption = styles.add_style('CaptionStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = 'Times New Roman'
    caption.font.size = Pt(10)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(80, 80, 80)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(12)

    # Monospace style
    code = styles.add_style('CodeStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = 'Courier New'
    code.font.size = Pt(10)
    code.font.color.rgb = RGBColor(50, 50, 50)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(4)
    code.paragraph_format.left_indent = Inches(0.5)

    # 1. COVER PAGE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run("SKILL DEVELOPMENT PROJECT REPORT\n\n\n")
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("GOOGLE MAPS INSPIRED ROUTE VISUALIZER\n\n")
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 115, 232)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("An interactive simulation platform for graph traversal and weighted shortest path algorithms utilizing binary Min-Heaps and DOM-direct rendering optimizations.\n\n\n\n")
    run.font.size = Pt(12)
    run.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run("Submitted in partial fulfillment of the academic requirements for the degree of\n")
    run.font.size = Pt(11)
    run = p.add_run("Bachelor of Technology in Computer Science & Engineering\n\n")
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    run = p.add_run("Developed By:\n")
    run.font.size = Pt(11)
    run = p.add_run("Student Name: Porandla Rajkumar\n")
    run.font.size = Pt(12)
    run.bold = True
    run = p.add_run("University Roll No: [YOUR-ROLL-NUMBER]\n\n\n")
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run("Department of Computer Science & Engineering\n")
    run.font.size = Pt(12)
    run.bold = True
    run = p.add_run("Affiliated University Name\n")
    run.font.size = Pt(11)
    run = p.add_run("Academic Session: 2025 - 2026\n")
    run.font.size = Pt(11)

    # 2. DECLARATION
    doc.add_page_break()
    h = doc.add_heading("DECLARATION", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "I, Porandla Rajkumar, student of Bachelor of Technology in Computer Science & Engineering, hereby declare "
        "that the project work entitled \"Google Maps Inspired Path Finding Visualizer\" submitted by me to the Department of "
        "Computer Science & Engineering is an authentic record of work carried out by me under academic supervision. "
        "This project has been developed as part of B.Tech 3rd Year Skill Development requirements."
    )
    doc.add_paragraph(
        "I declare that the software code, layout designs, and algorithms implementation presented in this report have been "
        "engineered by me, and that no part of this document has been plagiarized or submitted previously for any academic award."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("Date: 2026-07-09\nPlace: Hyderabad")
    r.font.size = Pt(11)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("_______________________\nPorandla Rajkumar")
    r2.font.size = Pt(11)
    r2.bold = True

    # 3. CERTIFICATE
    doc.add_page_break()
    h = doc.add_heading("CERTIFICATE OF APPROVAL", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This is to certify that the B.Tech Skill Development Project report entitled \"Google Maps Inspired Path Finding Visualizer\" "
        "is a record of work done by Porandla Rajkumar (Roll No: [YOUR-ROLL-NUMBER]) under our guidance. The work has been examined and "
        "approved for academic evaluation."
    )
    doc.add_paragraph(
        "This system represents a high-quality execution of graph search simulations incorporating direct-DOM paint pipelines to resolve "
        "browser layout throttling bottlenecks."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(96)
    r = p.add_run("Internal Guide\nAssistant Professor, CSE Dept.")
    r.font.size = Pt(11)
    r.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Head of the Department\nProfessor, CSE Dept.")
    r2.font.size = Pt(11)
    r2.bold = True

    # 4. ACKNOWLEDGEMENT
    doc.add_page_break()
    h = doc.add_heading("ACKNOWLEDGEMENT", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "I express my gratitude to the college management and the Head of Department for providing the computational tools "
        "and sandbox testing environments necessary to build this visualizer project."
    )
    doc.add_paragraph(
        "I am highly indebted to my internal guide, whose constructive guidance helped in refining the optimization techniques "
        "(such as CSS containment and direct-DOM rendering) to maintain a steady 60 FPS frame rate."
    )
    doc.add_paragraph(
        "Special thanks to my friends for their support in testing the keyboard shortcuts and JSON serialization modules, and to my parents "
        "for their encouragement."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("Porandla Rajkumar\nB.Tech CSE, 3rd Year")
    r.bold = True

    # 5. TABLE OF CONTENTS
    doc.add_page_break()
    h = doc.add_heading("TABLE OF CONTENTS", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contents = [
        ("Cover Page", "i"),
        ("Declaration", "ii"),
        ("Certificate of Approval", "iii"),
        ("Acknowledgement", "iv"),
        ("Abstract", "vi"),
        ("Chapter 1: Introduction", "1"),
        ("   1.1 Background & Graph Concepts", "1"),
        ("   1.2 Problem Statement", "2"),
        ("   1.3 Project Motivation", "2"),
        ("   1.4 Engineering Objectives", "3"),
        ("   1.5 System Scope", "3"),
        ("   1.6 Existing vs Proposed Solutions", "4"),
        ("Chapter 2: Academic Program Details", "6"),
        ("Chapter 3: Technical Architecture & Stack Selection", "8"),
        ("Chapter 4: System Design and Data Flows", "11"),
        ("Chapter 5: Traversal and Shortest Path Algorithms", "14"),
        ("   5.1 Breadth-First Search (BFS)", "14"),
        ("   5.2 Depth-First Search (DFS)", "16"),
        ("   5.3 Dijkstra's Shortest Path", "18"),
        ("   5.4 A* Heuristic Search", "20"),
        ("   5.5 Comparative Table", "23"),
        ("Chapter 6: Grid Maze and Obstacle Generation", "25"),
        ("   6.1 Recursive Division Maze", "25"),
        ("   6.2 Random Wall/Weight Generators", "26"),
        ("Chapter 7: UI & Core Features", "28"),
        ("Chapter 8: Module Implementation", "32"),
        ("Chapter 9: Verification and Testing Strategy", "36"),
        ("Chapter 10: Performance Challenges & Resolutions", "39"),
        ("   10.1 React Reconciliation Bottleneck", "39"),
        ("   10.2 Strict TypeScript Module Imports", "40"),
        ("Chapter 11: Future Scope", "42"),
        ("Chapter 12: Conclusion", "44"),
        ("References & Citations", "45"),
        ("Appendix: Source Code & Glossary", "46"),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Chapter / Topic"
    hdr_cells[1].text = "Page Number"
    set_cell_shading(hdr_cells[0], "E8F0FE")
    set_cell_shading(hdr_cells[1], "E8F0FE")
    
    for topic, page in contents:
        row_cells = table.add_row().cells
        row_cells[0].text = topic
        row_cells[1].text = page
        
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # 6. ABSTRACT
    doc.add_page_break()
    h = doc.add_heading("ABSTRACT", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Efficient pathfinding algorithms are vital for geographical information systems (GIS), network routing, and robotics. "
        "This project, \"Google Maps Inspired Path Finding Visualizer\", represents a high-performance simulation tool designed to model and evaluate "
        "four major graph traversal algorithms: Breadth-First Search (BFS), Depth-First Search (DFS), Dijkstra's Algorithm, and A* Search. "
        "Drawing inspiration from Google Maps' UI layout, the visualizer maps start coordinates (blue navigation arrow) and end coordinates (red pin), "
        "allowing users to paint roadblock wall barriers and traffic congestion delay weights (5x travel cost) directly on a grid canvas."
    )
    doc.add_paragraph(
        "To achieve fluid 60 FPS animations on large grids (such as 60x40), the visualizer implements direct-DOM paint injections during the playback phase, "
        "bypassing React state updates and virtual DOM diffing completely until the search sequence completes. Dijkstra and A* algorithms run on "
        "an optimized binary Min-Heap priority queue. Features like dynamic grid resizing, comparison dashboards, JSON map config file imports/exports, "
        "an onboarding tutorial slider, and an educational Viva inspector mode make this project a solid, production-ready portfolio piece."
    )

    # CHAPTER 1
    doc.add_page_break()
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    
    doc.add_heading("1.1 Background & Graph Concepts", level=2)
    doc.add_paragraph(
        "Graphs are mathematical structures used to model pairwise relations between objects, consisting of vertices (nodes) "
        "connected by edges. In routing systems, nodes represent physical locations (intersections or GPS coordinates) while edges represent roads. "
        "Pathfinding is the process of identifying the optimal path between a start node and an end node. Edges can be unweighted (all edges carry "
        "equal cost) or weighted (edges represent varying road lengths or traffic congestion delays)."
    )
    doc.add_paragraph(
        "Visualizing these traversals is an effective way to study computational complexity and data structures. For instance, Dijkstra's algorithm "
        "operates as a uniform-cost radial search, whereas A* guides the search direction towards the destination using heuristics, dramatically "
        "reducing evaluation overhead."
    )

    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "Most open-source pathfinding simulators suffer from rendering bottlenecks. When updating the search grid cell-by-cell at fast speeds, "
        "triggering React state modifications on every cell visited forces a re-render of the entire grid coordinate tree. For a 60x40 grid, "
        "rendering 2,400 cells results in severe lag, dropping browser frame rates below 15 FPS. Furthermore, standard simulators lack support for "
        "varying edge costs (like traffic delays) alongside walls, failing to demonstrate the difference between Dijkstra and unweighted BFS."
    )

    doc.add_heading("1.3 Project Motivation", level=2)
    doc.add_paragraph(
        "The motivation behind this project is to build a high-performance web platform that makes graph traversals easy to study. "
        "By styling the UI like Google Maps, users instantly connect academic complexities with real-world geographical systems, "
        "gaining an intuitive understanding of heuristics, priority queues, and rendering boundaries."
    )

    doc.add_heading("1.4 Engineering Objectives", level=2)
    doc.add_paragraph("The visualizer aims to achieve the following engineering goals:")
    doc.add_paragraph("• Design a fluid 60 FPS interactive grid that scales responsively across desktop and mobile displays.")
    doc.add_paragraph("• Implement BFS, DFS, Dijkstra, and A* Search algorithms using a custom-written binary Min-Heap priority queue.")
    doc.add_paragraph("• Build a comparative dashboard to benchmark speed, path cost, visited node counts, and memory footprint in real-time.")
    doc.add_paragraph("• Integrate a step-by-step playback controller (Play, Pause, Resume, Stop, Restart) that suspends/resumes animations without recalculating the route.")
    doc.add_paragraph("• Add a Viva Mode panel that visualizes the contents of the stack, queue, or priority heap at each step.")

    doc.add_heading("1.5 System Scope", level=2)
    doc.add_paragraph(
        "The application operates entirely on the client side using Vite + React + TypeScript, ensuring fast offline operations and lightweight "
        "deployment. It supports grid sizes from 20x20 up to 60x40 and features local storage preferences for map themes (dark vs light skins)."
    )

    doc.add_heading("1.6 Existing vs Proposed Solutions", level=2)
    doc.add_paragraph(
        "Existing simulators rely on state updates for each animation step. The proposed visualizer solves this by directly painting "
        "class names onto the DOM nodes during playback. React state is updated only when the path is fully drawn, ensuring 60 FPS rendering "
        "regardless of grid size."
    )

    # CHAPTER 2
    doc.add_page_break()
    doc.add_heading("CHAPTER 2: ACADEMIC PROGRAM DETAILS", level=1)
    doc.add_paragraph(
        "The Skill Development Project is designed to build practical engineering capabilities during B.Tech. "
        "Developing the G-Maps Inspired Route Visualizer provided hands-on experience in frontend performance tuning, "
        "TypeScript static compiler options, state machines, and data structures design. This project models standard professional "
        "software development lifecycle (SDLC) stages, from architectural planning to unit testing and build verification."
    )

    # CHAPTER 3
    doc.add_page_break()
    doc.add_heading("CHAPTER 3: TECHNICAL ARCHITECTURE & STACK SELECTION", level=1)
    
    doc.add_heading("3.1 React & Component Encapsulation", level=2)
    doc.add_paragraph(
        "React was selected for its modular component architecture. UI panels (Header, ControlPanel, StatsPanel, InfoPanel, VivaPanel, Grid) "
        "are encapsulated, making the codebase clean and maintainable. React hooks manage complex playback states."
    )

    doc.add_heading("3.2 TypeScript & Static Typing", level=2)
    doc.add_paragraph(
        "TypeScript catches bugs during compilation, before production build. Interfaces for nodes, coordinates, speed states, and telemetry data "
        "ensure strict type safety and eliminate runtime exceptions."
    )

    doc.add_heading("3.3 TailwindCSS v4", level=2)
    doc.add_paragraph(
        "TailwindCSS v4 is used for styling. Its modern CSS variables system supports light/dark theme toggles, and responsive classes "
        "ensure the UI scales fluidly on mobile screens."
    )

    doc.add_heading("3.4 Vite Build Tool", level=2)
    doc.add_paragraph(
        "Vite serves as the build tool, utilizing native ES modules (ESM) to provide instantaneous HMR and extremely fast build packaging."
    )

    doc.add_heading("3.5 Framer Motion", level=2)
    doc.add_paragraph(
        "Framer Motion drives smooth, physics-based slide transitions for the tutorial overlay guide and warning dialogs."
    )

    # CHAPTER 4
    doc.add_page_break()
    doc.add_heading("CHAPTER 4: SYSTEM DESIGN AND DATA FLOWS", level=1)
    doc.add_paragraph(
        "The application architecture uses a single-directional data flow. App.tsx coordinates parent state, while other components "
        "act as modular display units."
    )
    doc.add_paragraph(
        "[Figure 4.1: Component Layout and Communications Flow DFD]\n"
        "Description: Illustrates data flow from ControlPanel inputs through the parent App.tsx down to the Grid canvas and Viva inspect panel."
    ).style = 'CaptionStyle'

    # CHAPTER 5
    doc.add_page_break()
    doc.add_heading("CHAPTER 5: TRAVERSAL AND SHORTEST PATH ALGORITHMS", level=1)
    
    # 5.1 BFS
    doc.add_heading("5.1 Breadth-First Search (BFS)", level=2)
    doc.add_paragraph(
        "BFS explores nodes level-by-level, utilizing a FIFO Queue. It guarantees the shortest path on unweighted graphs."
    )
    doc.add_paragraph("Complexity: O(V + E) Time | O(V) Space.")

    # 5.2 DFS
    doc.add_heading("5.2 Depth-First Search (DFS)", level=2)
    doc.add_paragraph(
        "DFS explores branches as deep as possible before backtracking, using a LIFO Stack. It does NOT guarantee the shortest path."
    )
    doc.add_paragraph("Complexity: O(V + E) Time | O(V) Space.")

    # 5.3 Dijkstra
    doc.add_heading("5.3 Dijkstra's Shortest Path", level=2)
    doc.add_paragraph(
        "Dijkstra solves the single-source shortest path problem on weighted graphs. It extracts the lowest-cost unvisited nodes using a "
        "binary Min-Heap priority queue, adapting routes around heavy traffic congestion blocks."
    )
    doc.add_paragraph("Complexity: O((V + E) log V) Time | O(V) Space.")

    # 5.4 A* Search
    doc.add_heading("5.4 A* Heuristic Search", level=2)
    doc.add_paragraph(
        "A* Search guides the search toward the destination by minimizing f(n) = g(n) + h(n), where h(n) represents the Manhattan distance "
        "to the end pin. This reduces search overhead compared to Dijkstra."
    )
    doc.add_paragraph("Complexity: O(E log V) Average Time | O(b^d) Space.")

    # 5.5 Comparison
    doc.add_heading("5.5 Comparative Benchmark Layout", level=2)
    
    table = doc.add_table(rows=1, cols=6)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Algorithm"
    hdr_cells[1].text = "Data Structure"
    hdr_cells[2].text = "Weighted?"
    hdr_cells[3].text = "Shortest Path?"
    hdr_cells[4].text = "Time Complexity"
    hdr_cells[5].text = "Space Complexity"
    
    for cell in hdr_cells:
        set_cell_shading(cell, "E8F0FE")

    runs = [
        ("A* Search", "Min-Heap", "Yes", "Yes", "O(E log V) Avg", "O(b^d)"),
        ("Dijkstra", "Min-Heap", "Yes", "Yes", "O((V+E) log V)", "O(V)"),
        ("BFS", "Queue", "No", "Yes (Unweighted)", "O(V+E)", "O(V)"),
        ("DFS", "Stack", "No", "No", "O(V+E)", "O(V)"),
    ]
    
    for algo, ds, w, sp, t, s in runs:
        row_cells = table.add_row().cells
        row_cells[0].text = algo
        row_cells[1].text = ds
        row_cells[2].text = w
        row_cells[3].text = sp
        row_cells[4].text = t
        row_cells[5].text = s

    p_cap = doc.add_paragraph("Table 5.1: Performance and complexity profiles of the four traversal algorithms.")
    p_cap.style = 'CaptionStyle'

    # CHAPTER 6
    doc.add_page_break()
    doc.add_heading("CHAPTER 6: GRID MAZE AND OBSTACLE GENERATION", level=1)
    doc.add_heading("6.1 Recursive Division Maze", level=2)
    doc.add_paragraph(
        "Recursive Division divides the grid using horizontal or vertical walls, leaves a single grid block passage open, "
        "and recursively divides the newly created subgrids. Wall alignments are placed on even indices and passages on odd indices, "
        "generating consistent, structured street layouts."
    )
    doc.add_heading("6.2 Random Wall/Weight Generators", level=2)
    doc.add_paragraph(
        "Random Roadblocks places solid walls with a 30% density. Random Traffic Jams distributes weighted cells "
        "(5x cost) across 25% of the grid, allowing users to test how weighted pathfinders adapt routing."
    )

    # CHAPTER 7
    doc.add_page_break()
    doc.add_heading("CHAPTER 7: UI & CORE FEATURES", level=1)
    
    features = [
        ("Interactive Grid Canvas", "Supports dynamic sizing options: 20x20, 25x25, 30x30, 40x25, 50x30, and 60x40. Grid is fully responsive."),
        ("Drag and Drop Pins", "Allows real-time repositioning of start (current location) and destination pin markers."),
        ("Grid Marker Tools", "Draw walls (solid roadblocks) or weighted cells (traffic delays) dynamically by clicking and dragging."),
        ("Dark/Light Mode Skins", "Clean map skins inspired by Google Maps, configured using Tailwind CSS variables."),
        ("Playback controller", "Full Play, Pause, Resume, Stop, and Restart controls that suspend/resume execution index arrays instantly."),
        ("Comparison Benchmarking", "Runs all algorithms simultaneously on the same map to log execution time, visited counts, path cost, and memory estimate. Exports to JSON/CSV."),
        ("Onboarding Guide", "Step-by-step slideshow walkthrough built with Framer Motion, introducing visualizer elements."),
        ("Save & Load Maps", "Allows saving and restoring roadblock and traffic layouts as local JSON files."),
        ("Viva Inspector Mode", "Inspects active Queue, Stack, or Heap contents at each step, making it ideal for university demonstrations."),
        ("Keyboard Hotkey Shortcuts", "Global hotkeys (Space to play, P to pause, R to reset, H to open guide, etc.) for quick operations.")
    ]
    for title, desc in features:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # CHAPTER 8
    doc.add_page_break()
    doc.add_heading("CHAPTER 8: MODULE IMPLEMENTATION", level=1)
    
    modules = [
        ("Header Component", "Contains branding logo, active calculation pulse alerts, and theme toggle buttons."),
        ("Grid Component", "Houses the grid matrix, managing click/drag events and bubbling up keyboard arrow directions."),
        ("Node Component", "Individual grid cell. Memoized via React.memo and styled with CSS containment ('contain: layout style paint') to optimize frame rates."),
        ("Control Panel Component", "Main dashboard containing algorithm selector, speed select, grid size choice, JSON triggers, and playback buttons."),
        ("Statistics Panel Component", "Displays live telemetry stats: CPU compute time, frames per second, wall count, and memory footprint."),
        ("Comparison Dashboard", "Benchmarks all pathfinders, displaying comparison table metrics and highlight badges."),
        ("Viva Panel Component", "Dynamic inspector updating data structure snapshots at each animation step."),
        ("Toast Component", "Status notifications panel rendering success/warning/error alerts in the corner."),
        ("Algorithms Module", "Implements Dijkstra, A*, BFS, DFS, and Heap classes under strict verbatim type checking.")
    ]
    for title, desc in modules:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # CHAPTER 9
    doc.add_page_break()
    doc.add_heading("CHAPTER 9: VERIFICATION AND TESTING STRATEGY", level=1)
    doc.add_paragraph(
        "A dedicated test suite in `src/visualizer.test.ts` validates the correctness of the traversal algorithms. "
        "Tests verify grid boundaries, BFS unweighted routing lengths, and Dijkstra's routing accuracy with walls. "
        "A production bundle test compiles TypeScript and builds package distributions without any errors or warnings."
    )

    # CHAPTER 10
    doc.add_page_break()
    doc.add_heading("CHAPTER 10: PERFORMANCE CHALLENGES & RESOLUTIONS", level=1)
    
    doc.add_heading("10.1 React Reconciliation Bottleneck", level=2)
    doc.add_paragraph(
        "Problem: When rendering pathfinding animations, updating the React grid state at 10ms intervals caused the entire "
        "Grid component to re-render. On a 50x30 grid, this triggered 1,500 state updates, lagging frame rates below 15 FPS."
    )
    doc.add_paragraph(
        "Resolution: Implemented direct-DOM injection during animation playback. The script targets cell IDs directly via "
        "`document.getElementById` to add/remove classes like `.node-visited`, avoiding React state updates entirely until the route is fully drawn. "
        "This maintains a steady 60 FPS."
    )

    doc.add_heading("10.2 Strict TypeScript Module Imports", level=2)
    doc.add_paragraph(
        "Problem: TypeScript's strict module syntax requires explicit `import type` annotations for pure interface files. "
        "Mixing type and value imports threw compiler errors during `npm run build`."
    )
    doc.add_paragraph(
        "Resolution: Refactored all data interfaces to use `import type { GridNode }` imports, separating value imports (like constants and functions) "
        "from types."
    )

    # CHAPTER 11
    doc.add_page_break()
    doc.add_heading("CHAPTER 11: FUTURE SCOPE", level=1)
    
    enhancements = [
        ("AI Route Prediction", "Train neural networks to predict shortest routing coordinates by analyzing historical traffic congestion states."),
        ("Real Maps API Integration", "Load actual coordinates from OpenStreetMap API, allowing users to visual real-world pathfinding routes."),
        ("Multi-Destination Support", "Enable multi-stop routing to resolve Traveling Salesperson Problem (TSP) scenarios on graphs."),
        ("3D Route Visualizations", "Utilize Three.js to render elevated terrain, bridges, and flyovers in an interactive 3D map space.")
    ]
    for title, desc in enhancements:
        doc.add_heading(title, level=2)
        doc.add_paragraph(desc)

    # CHAPTER 12
    doc.add_page_break()
    doc.add_heading("CHAPTER 12: CONCLUSION", level=1)
    doc.add_paragraph(
        "The \"Google Maps Inspired Path Finding Visualizer\" is a complete Skill Development Project. "
        "By implementing four graph search algorithms (BFS, DFS, Dijkstra, A*) guided by custom-written priority queue Min-Heaps, the application "
        "provides a clear, high-performance representation of pathfinding theory. Direct-DOM optimizations resolve browser reflow bottlenecks, "
        "guaranteeing a steady 60 FPS. With features like saving/loading maps, comparing algorithms, and inspecting stack/queue states in Viva Mode, "
        "the project stands as a portfolio-quality software asset ready for university evaluation and production deployment."
    )

    # REFERENCES
    doc.add_page_break()
    h = doc.add_heading("REFERENCES", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    refs = [
        "1. Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press.",
        "2. Russell, S., & Norvig, P. (2020). Artificial Intelligence: A Modern Approach (4th ed.). Pearson.",
        "3. React Official Documentation. Component lifecycle, Hooks, and State management. https://react.dev",
        "4. MDN Web Docs. CSS Containment (layout, style, paint) for rendering optimization. https://developer.mozilla.org"
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # APPENDIX
    doc.add_page_break()
    h = doc.add_heading("APPENDIX", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading("A. Min-Heap Class Snippet", level=2)
    p_code = doc.add_paragraph()
    p_code.style = 'CodeStyle'
    p_code.add_run(
        "export class MinHeap {\n"
        "  private heap: GridNode[] = [];\n"
        "  public insert(node: GridNode): void {\n"
        "    this.heap.push(node);\n"
        "    this.heapifyUp(this.heap.length - 1);\n"
        "  }\n"
        "  public extractMin(): GridNode | null {\n"
        "    if (this.isEmpty()) return null;\n"
        "    const min = this.heap[0];\n"
        "    const end = this.heap.pop();\n"
        "    if (this.heap.length > 0 && end !== undefined) {\n"
        "      this.heap[0] = end;\n"
        "      this.heapifyDown(0);\n"
        "    }\n"
        "    return min;\n"
        "  }\n"
        "}"
    )

    doc.save("G-Maps_Route_Visualizer_Report.docx")
    print("Polished Report v2 generated successfully as G-Maps_Route_Visualizer_Report.docx")

if __name__ == '__main__':
    create_report_v2()
